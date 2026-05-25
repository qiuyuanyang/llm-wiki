#!/usr/bin/env python3
# SPDX-License-Identifier: MIT
"""
wiki_watcher.py — LLM Wiki automation daemon

Watches the raw/ directory for new files, calls the configured LLM via the
Anthropic SDK or an OpenAI-compatible endpoint, and writes the results to wiki/.

Dependencies:
    pip install anthropic watchdog rich

Usage:
    python scripts/wiki_watcher.py                    # foreground mode
    python scripts/wiki_watcher.py --daemon           # background daemon
    python scripts/wiki_watcher.py --status           # show running status
    python scripts/wiki_watcher.py --stop             # stop daemon
"""

import os
import re
import sys
import json
import time
import signal
import hashlib
import argparse
import threading
import subprocess
from pathlib import Path
from datetime import datetime
from queue import Queue, Empty

# Ensure project root is importable for vector_ingest
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

# ── Dependency check ──────────────────────────────────────────────────────────
try:
    from watchdog.observers import Observer
    from watchdog.events import FileSystemEventHandler
    from rich.console import Console
    from rich.panel import Panel
    from rich.live import Live
    from rich.table import Table
    from rich import print as rprint
except ImportError as e:
    print(f"Missing dependency: {e}")
    print("Run: pip install anthropic watchdog rich")
    sys.exit(1)

console = Console()

# ── Configuration ─────────────────────────────────────────────────────────────
SUPPORTED_EXTENSIONS = {'.md', '.txt', '.pdf', '.html', '.rst', '.epub', '.doc', '.docx', '.xls', '.xlsx'}
IGNORE_PATTERNS = {'*.tmp', '*.DS_Store', '.gitkeep', '*.swp', '._*'}
DEBOUNCE_SECONDS = 3       # wait for file write to complete before processing
MAX_CONCURRENT_JOBS = 2    # maximum parallel ingest jobs
AUTO_HOTSPOT_THRESHOLD = 5 # trigger hotspot analysis every N ingests


class WikiConfig:
    def __init__(self, wiki_root: Path):
        self.wiki_root = wiki_root
        self.raw_dir = wiki_root / "raw" / "sources"
        self.wiki_dir = wiki_root / "wiki"
        self.claude_md = wiki_root / "LLM.md"
        self.index_md = wiki_root / "wiki" / "index.md"
        self.log_md = wiki_root / "wiki" / "log.md"
        self.state_file = wiki_root / ".wiki_watcher_state.json"
        self.pid_file = wiki_root / ".wiki_watcher.pid"

        for d in [self.raw_dir, self.wiki_dir]:
            d.mkdir(parents=True, exist_ok=True)

    def load_state(self) -> dict:
        if self.state_file.exists():
            return json.loads(self.state_file.read_text())
        return {"processed": {}, "ingest_count": 0, "queue": []}

    def save_state(self, state: dict):
        self.state_file.write_text(json.dumps(state, indent=2, ensure_ascii=False))

    def load_schema(self) -> str:
        if self.claude_md.exists():
            return self.claude_md.read_text(encoding='utf-8')
        return ""


# ── Model routing layer ───────────────────────────────────────────────────────
#
# wiki_config.json example:
# {
#   "active": {"provider": "anthropic", "model": "claude-opus-4-6"},
#   "providers": {
#     "anthropic": {"api_key_env": "ANTHROPIC_API_KEY"},
#     "local":     {"base_url": "http://localhost:11434/v1", "api_key": "ollama"},
#     "dashscope": {"base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
#                   "api_key_env": "DASHSCOPE_API_KEY"}
#   }
# }
#
# Switch model: python wiki_watcher.py --model local,Qwen3.6-35B-A3B-FP8
#               ./wiki.sh model dashscope,qwen-max
#
class ModelRouter:
    DEFAULT_CONFIG = {
        "active": {"provider": "anthropic", "model": "claude-opus-4-6"},
        "providers": {
            "anthropic": {
                "api_key_env": "ANTHROPIC_API_KEY"
            },
            "local": {
                "base_url": "http://localhost:11434/v1",
                "api_key": "ollama"
            },
            "dashscope": {
                "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
                "api_key_env": "DASHSCOPE_API_KEY"
            },
            "openai": {
                "api_key_env": "OPENAI_API_KEY"
            }
        }
    }

    def __init__(self, wiki_root: Path):
        self.config_path = wiki_root / "wiki_config.json"
        self.cfg = self._load()

    @staticmethod
    def _expand_env(obj):
        """Recursively expand ${VAR_NAME} placeholders to their environment values.
        Undefined variables are left as-is so missing config is easy to diagnose."""
        if isinstance(obj, str):
            return re.sub(r'\$\{([^}]+)\}',
                          lambda m: os.environ.get(m.group(1), m.group(0)),
                          obj)
        if isinstance(obj, dict):
            return {k: ModelRouter._expand_env(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [ModelRouter._expand_env(i) for i in obj]
        return obj

    def _load(self) -> dict:
        if self.config_path.exists():
            data = json.loads(self.config_path.read_text(encoding='utf-8'))
            # Merge missing provider defaults while preserving user customisations
            default_providers = self.DEFAULT_CONFIG["providers"]
            data.setdefault("providers", {})
            for name, defaults in default_providers.items():
                data["providers"].setdefault(name, defaults)
            # _raw retains original ${VAR} placeholders for disk writes
            # (never persist plaintext secrets)
            self._raw = json.loads(json.dumps(data))
            return self._expand_env(data)
        # First run: write default config
        self.config_path.write_text(
            json.dumps(self.DEFAULT_CONFIG, indent=2, ensure_ascii=False)
        )
        self._raw = json.loads(json.dumps(self.DEFAULT_CONFIG))
        return dict(self.DEFAULT_CONFIG)

    def _save_raw(self):
        """Write the raw (placeholder-containing) config back to disk — never the expanded plaintext."""
        self.config_path.write_text(
            json.dumps(self._raw, indent=2, ensure_ascii=False)
        )

    def set_model(self, spec: str):
        """Switch model. spec format: 'provider,model_name' or 'model_name' (defaults to anthropic)."""
        if ',' in spec:
            provider, model = spec.split(',', 1)
        else:
            provider, model = 'anthropic', spec
        provider, model = provider.strip(), model.strip()
        self.cfg['active'] = {'provider': provider, 'model': model}
        self._raw['active'] = {'provider': provider, 'model': model}
        self._save_raw()
        console.print(f"[green]Model switched: [bold]{provider}[/bold] / [bold]{model}[/bold][/green]")

    @property
    def active_provider(self) -> str:
        return self.cfg['active']['provider']

    @property
    def active_model(self) -> str:
        return self.cfg['active']['model']

    def call(self, system: str, user: str, max_tokens: int) -> str:
        """Unified call interface — dispatches to the active provider's SDK."""
        provider = self.active_provider
        model = self.active_model
        console.print(f"[dim]  → {provider} / {model}[/dim]")

        if provider == 'anthropic':
            return self._call_anthropic(model, system, user, max_tokens)
        else:
            return self._call_openai_compat(provider, model, system, user, max_tokens)

    def _call_anthropic(self, model: str, system: str, user: str, max_tokens: int) -> str:
        try:
            import anthropic as _anthropic
        except ImportError:
            raise RuntimeError("Anthropic provider requires: pip install anthropic")
        client = _anthropic.Anthropic(
            api_key=os.environ.get(
                self.cfg['providers'].get('anthropic', {}).get('api_key_env', 'ANTHROPIC_API_KEY')
            )
        )
        resp = client.messages.create(
            model=model,
            max_tokens=max_tokens,
            system=system,
            messages=[{"role": "user", "content": user}]
        )
        return resp.content[0].text

    def _call_openai_compat(self, provider: str, model: str,
                             system: str, user: str, max_tokens: int) -> str:
        try:
            import openai as _openai
        except ImportError:
            raise RuntimeError(f"{provider} provider requires: pip install openai")
        pconf = self.cfg['providers'].get(provider, {})
        base_url = pconf.get('base_url')
        api_key_env = pconf.get('api_key_env', '')
        api_key = pconf.get('api_key') or os.environ.get(api_key_env, 'none')

        # Normalise base_url: if the user supplied a full endpoint (ending in
        # /chat/completions), strip it back to /v1 — the SDK appends the path itself.
        if base_url and base_url.rstrip('/').endswith('/chat/completions'):
            base_url = base_url.rstrip('/')[:-len('/chat/completions')]

        # Force NO_PROXY for local/intranet addresses to bypass system proxies
        # (e.g. a transparent Squid proxy on macOS that intercepts 127.x traffic).
        if base_url:
            from urllib.parse import urlparse
            host = urlparse(base_url).hostname or ''
            _local_prefixes = ('127.', '10.', '192.168.', '172.', 'localhost', '::1')
            if any(host == p or host.startswith(p) for p in _local_prefixes):
                existing = os.environ.get('NO_PROXY', '')
                if host not in existing:
                    os.environ['NO_PROXY'] = f"{existing},{host},localhost,127.0.0.1".lstrip(',')
                    os.environ['no_proxy'] = os.environ['NO_PROXY']

        kwargs = {"api_key": api_key, "max_retries": 2}
        if base_url:
            kwargs["base_url"] = base_url

        client = _openai.OpenAI(**kwargs)
        resp = client.chat.completions.create(
            model=model,
            max_tokens=max_tokens,
            messages=[
                {"role": "system", "content": system},
                {"role": "user",   "content": user}
            ]
        )
        return resp.choices[0].message.content


# ── LLM call layer ────────────────────────────────────────────────────────────
class WikiAgent:
    def __init__(self, config: WikiConfig):
        self.config = config
        self.router = ModelRouter(config.wiki_root)
        self.schema = config.load_schema()

    def _read_file_content(self, file_path: Path) -> str:
        """Read file contents, handling different formats."""
        suffix = file_path.suffix.lower()

        if suffix == '.pdf':
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    return "\n\n".join(p.extract_text() or "" for p in pdf.pages)
            except ImportError:
                return (
                    f"[PDF file: {file_path.name}]\n"
                    "Note: install pdfplumber to extract PDF text automatically\n"
                    "`pip install pdfplumber`"
                )

        if suffix in ('.docx',):
            try:
                from docx import Document
                doc = Document(file_path)
                parts = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        parts.append(p.text.strip())
                for table in doc.tables:
                    parts.append(self._format_table(table.rows))
                return "\n\n".join(parts)
            except ImportError:
                return f"[DOCX file: {file_path.name}]\nNote: install python-docx: `pip install python-docx`"

        if suffix in ('.doc',):
            # Legacy .doc: try antiword or textract fallback
            try:
                import subprocess
                result = subprocess.run(
                    ['antiword', str(file_path)],
                    capture_output=True, text=True, timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass
            return (
                f"[DOC file: {file_path.name}]\n"
                "注意: .doc 格式需要安装 antiword (apt install antiword) 才能提取文本。"
                "建议转换为 .docx 格式。"
            )

        if suffix in ('.xlsx', '.xls'):
            try:
                if suffix == '.xlsx':
                    import openpyxl
                    wb = openpyxl.load_workbook(file_path, data_only=True)
                else:
                    import xlrd
                    wb = xlrd.open_workbook(file_path)
                    return self._read_xls_book(wb)

                parts = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    if hasattr(ws, 'max_row'):  # openpyxl
                        rows = list(ws.iter_rows(values_only=True))
                    else:
                        rows = []
                    parts.append(f"## 工作表: {sheet_name}")
                    parts.append(self._format_rows(rows))
                return "\n\n".join(parts)
            except ImportError:
                return f"[{suffix.upper()} file: {file_path.name}]\nNote: install openpyxl/xlrd: `pip install openpyxl xlrd`"

        try:
            return file_path.read_text(encoding='utf-8')
        except UnicodeDecodeError:
            return file_path.read_text(encoding='latin-1')

    @staticmethod
    def _format_rows(rows) -> str:
        """Format rows as a Markdown table."""
        if not rows:
            return '_(空工作表)_'
        # Filter fully empty rows
        rows = [r for r in rows if any(c is not None and str(c).strip() for c in r)]
        if not rows:
            return '_(空工作表)_'

        # Normalize
        max_cols = max(len(r) for r in rows)
        normalized = []
        for r in rows:
            nr = [str(c) if c is not None else '' for c in r]
            while len(nr) < max_cols:
                nr.append('')
            normalized.append(nr)

        lines = []
        lines.append('| ' + ' | '.join(normalized[0]) + ' |')
        lines.append('| ' + ' | '.join('---' for _ in normalized[0]) + ' |')
        for r in normalized[1:]:
            lines.append('| ' + ' | '.join(r) + ' |')
        return '\n'.join(lines)

    @staticmethod
    def _format_table(rows) -> str:
        """Format docx Table rows as Markdown table."""
        if not rows:
            return ''
        cells_list = []
        for row in rows:
            cells = [cell.text.strip() for cell in row.cells]
            cells_list.append(cells)

        max_cols = max(len(c) for c in cells_list) if cells_list else 0
        for cells in cells_list:
            while len(cells) < max_cols:
                cells.append('')

        lines = []
        lines.append('| ' + ' | '.join(cells_list[0]) + ' |')
        lines.append('| ' + ' | '.join('---' for _ in cells_list[0]) + ' |')
        for cells in cells_list[1:]:
            lines.append('| ' + ' | '.join(cells) + ' |')
        return '\n'.join(lines)

    @staticmethod
    def _read_xls_book(wb) -> str:
        """Read xlrd workbook to text."""
        parts = []
        for si in range(wb.nsheets):
            ws = wb.sheet_by_index(si)
            parts.append(f"## 工作表: {ws.name}")
            rows = []
            for r in range(ws.nrows):
                rows.append([ws.cell_value(r, c) for c in range(ws.ncols)])
            parts.append(WikiAgent._format_rows(rows))
        return "\n\n".join(parts)

    def _read_wiki_context(self) -> str:
        """Read the current wiki state (index + recent log entries)."""
        parts = []

        index_path = self.config.index_md
        if index_path.exists():
            parts.append(f"=== wiki/index.md ===\n{index_path.read_text(encoding='utf-8')}")

        log_path = self.config.log_md
        if log_path.exists():
            log_content = log_path.read_text(encoding='utf-8')
            lines = log_content.split('\n')
            recent_entries = []
            entry_count = 0
            for line in reversed(lines):
                recent_entries.insert(0, line)
                if line.startswith('## ['):
                    entry_count += 1
                    if entry_count >= 20:
                        break
            parts.append("=== wiki/log.md (recent entries) ===\n" + '\n'.join(recent_entries))

        return "\n\n".join(parts)

    def ingest_file(self, file_path: Path) -> dict:
        """Process a single new file and return the structured result."""
        console.print(f"[cyan]Processing: {file_path.name}[/cyan]")

        file_content = self._read_file_content(file_path)
        wiki_context = self._read_wiki_context()
        today = datetime.now().strftime('%Y-%m-%d')

        system_prompt = f"""你是 wiki 知识库的管理 Agent。

{self.schema}

---
当前任务：自动导入新的源材料
日期：{today}
文件：raw/sources/{file_path.name}

规则：
1. 你的回复必须是以下格式的合法 JSON — 不要有其他内容。
2. 所有 wiki 文件内容放在 "files" 数组中。
3. 文件路径相对于 wiki 根目录。
4. 你必须更新 wiki/index.md 和 wiki/log.md。
5. **所有生成的内容（标题、描述、总结、解释等）必须使用中文。** 源数据内容（代码、配置、命令、IP地址等）保持原文不变。

返回格式（严格 JSON，不要包裹其他文本）：
{{
  "summary": "一句话描述本次导入的核心发现",
  "files": [
    {{
      "path": "wiki/sources/xxx.md",
      "action": "create",
      "content": "包含前置元数据的完整文件内容"
    }},
    {{
      "path": "wiki/index.md",
      "action": "update",
      "content": "完整的更新后的 index.md 内容"
    }}
  ],
  "new_pages": ["wiki/sources/xxx.md", "wiki/entities/yyy.md"],
  "updated_pages": ["wiki/index.md", "wiki/log.md"],
  "discoveries": ["关键发现 1", "关键发现 2"],
  "follow_up_questions": ["此材料引发的新问题 1", "问题 2"]
}}"""

        user_message = f"""请处理以下源材料并更新 wiki：

=== 源材料 ===
{file_content[:8000]}

=== 当前 wiki 状态 ===
{wiki_context[:4000]}"""

        try:
            raw_text = self.router.call(system_prompt, user_message, max_tokens=8000).strip()

            # Strip Qwen3 chain-of-thought <think>...</think> blocks
            raw_text = re.sub(r'<think>.*?</think>', '', raw_text, flags=re.DOTALL).strip()

            # Strip optional markdown code-fence wrapping (```json ... ``` or ``` ... ```)
            if raw_text.startswith("```"):
                raw_text = raw_text.split('\n', 1)[1] if '\n' in raw_text else raw_text[3:]
                raw_text = raw_text.rsplit('```', 1)[0].strip()

            # Extract the first complete JSON object (ignore any preamble/postamble)
            brace_start = raw_text.find('{')
            brace_end = raw_text.rfind('}')
            if brace_start != -1 and brace_end != -1:
                raw_text = raw_text[brace_start:brace_end + 1]

            result = json.loads(raw_text)
            result["source_file"] = str(file_path)
            result["processed_at"] = datetime.now().isoformat()
            return result

        except json.JSONDecodeError as e:
            console.print(f"[red]JSON parse error: {e}[/red]")
            console.print(f"[dim]Raw response excerpt: {raw_text[:300]}[/dim]")
            try:
                import json_repair
                result = json_repair.loads(raw_text)
                if isinstance(result, dict) and result:
                    console.print("[yellow]JSON repaired via json_repair[/yellow]")
                    result["source_file"] = str(file_path)
                    result["processed_at"] = datetime.now().isoformat()
                    return result
            except ImportError:
                pass
            except Exception:
                pass
            return {"error": str(e), "raw": raw_text[:500]}
        except Exception as e:
            console.print(f"[red]API call error: {e}[/red]")
            return {"error": str(e)}

    def generate_hotspot_analysis(self) -> dict:
        """Generate a hotspot analysis and knowledge graph summary."""
        console.print("[magenta]Generating hotspot analysis...[/magenta]")

        wiki_context = self._read_wiki_context()
        today = datetime.now().strftime('%Y-%m-%d')

        all_pages = []
        for page in self.config.wiki_dir.rglob("*.md"):
            if page.name not in ('index.md', 'log.md', 'dashboard.md'):
                all_pages.append(str(page.relative_to(self.config.wiki_root)))

        system_prompt = f"""你是 wiki 知识库分析 Agent。分析当前知识库状态并生成热点摘要。

返回严格 JSON：
{{
  "hotspot_page": {{
    "path": "wiki/queries/{today}-hotspot-analysis.md",
    "content": "完整的热点分析 Markdown 页面内容（包含 YAML 前置元数据）"
  }},
  "overview_update": {{
    "path": "wiki/overview.md",
    "content": "完整的更新后的 overview.md 内容"
  }},
  "key_themes": ["主题 1", "主题 2", "主题 3"],
  "emerging_connections": ["新发现的概念关联 1", "关联 2"],
  "recommended_next_sources": ["建议下一步阅读的材料类型"]
}}"""

        user_message = f"""分析当前知识库：识别热点主题、核心概念关系和知识空白。

当前页面列表：
{chr(10).join(all_pages[:50])}

{wiki_context[:5000]}"""

        try:
            raw_text = self.router.call(system_prompt, user_message, max_tokens=6000).strip()
            raw_text = re.sub(r'<think>.*?</think>', '', raw_text, flags=re.DOTALL).strip()
            if raw_text.startswith("```"):
                raw_text = raw_text.split('\n', 1)[1].rsplit('```', 1)[0].strip()
            brace_start, brace_end = raw_text.find('{'), raw_text.rfind('}')
            if brace_start != -1 and brace_end != -1:
                raw_text = raw_text[brace_start:brace_end + 1]
            return json.loads(raw_text)
        except Exception as e:
            console.print(f"[red]Hotspot analysis error: {e}[/red]")
            return {"error": str(e)}


# ── File write layer ──────────────────────────────────────────────────────────
class WikiWriter:
    def __init__(self, config: WikiConfig):
        self.config = config

    def apply_result(self, result: dict) -> list:
        """Write Agent results to the filesystem."""
        written = []

        if "error" in result:
            console.print(f"[red]Processing failed: {result['error']}[/red]")
            return written

        for file_spec in result.get("files", []):
            file_path = self.config.wiki_root / file_spec["path"]
            file_path.parent.mkdir(parents=True, exist_ok=True)

            content = file_spec.get("content", "")

            # For log.md: append rather than overwrite
            if file_spec.get("action") == "update" and file_path.name == "log.md":
                existing = file_path.read_text(encoding='utf-8') if file_path.exists() else ""
                if not existing:
                    file_path.write_text(content, encoding='utf-8')
                elif content not in existing:
                    if '## [' in content:
                        new_entry = '## [' + content.split('## [')[-1]
                    else:
                        new_entry = content
                    file_path.write_text(existing.rstrip() + "\n\n" + new_entry, encoding='utf-8')
            else:
                file_path.write_text(content, encoding='utf-8')

            written.append(file_spec["path"])
            console.print(f"  [green]✓[/green] {file_spec['action']}: {file_spec['path']}")

        return written

    def apply_hotspot(self, result: dict):
        """Write hotspot analysis results to disk."""
        if "error" in result:
            return

        for key in ["hotspot_page", "overview_update"]:
            if key in result:
                spec = result[key]
                file_path = self.config.wiki_root / spec["path"]
                file_path.parent.mkdir(parents=True, exist_ok=True)
                file_path.write_text(spec["content"], encoding='utf-8')
                console.print(f"  [magenta]hotspot[/magenta] written: {spec['path']}")


# ── File watch layer ──────────────────────────────────────────────────────────
class RawDirectoryHandler(FileSystemEventHandler):
    def __init__(self, job_queue: Queue, config: WikiConfig):
        self.job_queue = job_queue
        self.config = config
        self.pending = {}  # debounce map: file path → scheduled process time
        self._lock = threading.Lock()

    def _should_process(self, file_path: str) -> bool:
        path = Path(file_path)
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            return False
        if path.name.startswith('._'):  # macOS resource fork files
            return False
        if any(path.match(p) for p in IGNORE_PATTERNS):
            return False
        if 'wiki' in path.parts:  # avoid processing wiki/ output (would loop)
            return False
        return True

    def _schedule_job(self, file_path: str):
        with self._lock:
            self.pending[file_path] = time.time() + DEBOUNCE_SECONDS

    def _debounce_worker(self):
        """Background thread: wait for file stability before enqueuing."""
        while True:
            time.sleep(0.5)
            now = time.time()
            with self._lock:
                ready = [p for p, t in self.pending.items() if now >= t]
                for p in ready:
                    del self.pending[p]
                    self.job_queue.put(("ingest", Path(p)))
                    console.print(f"[yellow]New file detected: {Path(p).name}[/yellow]")

    def start_debounce_worker(self):
        t = threading.Thread(target=self._debounce_worker, daemon=True)
        t.start()

    def on_created(self, event):
        if not event.is_directory and self._should_process(event.src_path):
            self._schedule_job(event.src_path)

    def on_moved(self, event):
        if not event.is_directory and self._should_process(event.dest_path):
            self._schedule_job(event.dest_path)


# ── Main controller ───────────────────────────────────────────────────────────
class WikiWatcher:
    def __init__(self, wiki_root: Path):
        self.config = WikiConfig(wiki_root)
        self.agent = WikiAgent(self.config)
        self.writer = WikiWriter(self.config)
        self.job_queue = Queue()
        self.state = self.config.load_state()
        self.running = False

    def _file_hash(self, path: Path) -> str:
        return hashlib.md5(path.read_bytes()).hexdigest()

    def _already_processed(self, path: Path) -> bool:
        file_hash = self._file_hash(path)
        return self.state["processed"].get(str(path)) == file_hash

    def _mark_processed(self, path: Path):
        self.state["processed"][str(path)] = self._file_hash(path)
        self.state["ingest_count"] = self.state.get("ingest_count", 0) + 1
        self.config.save_state(self.state)

    def _sync_vectors(self, written_files: list):
        """Re-sync vector store after wiki pages are written."""
        try:
            from scripts.vector_ingest import ingest_all
            from config import WikiConfig
            from embedding_client import EmbeddingClient

            config = WikiConfig()
            embedding_client = EmbeddingClient(config)
            result = ingest_all(config, embedding_client, show_progress=False)
            console.print(f"  [dim]📐 Vectors synced: {result['success']} pages indexed[/dim]")
        except Exception as e:
            console.print(f"  [yellow]⚠ Vector sync failed: {e}[/yellow]")

    def process_job(self, job_type: str, file_path: Path):
        if job_type == "ingest":
            if self._already_processed(file_path):
                console.print(f"[dim]Skipping (already processed): {file_path.name}[/dim]")
                return

            result = self.agent.ingest_file(file_path)
            written = self.writer.apply_result(result)

            if written:
                self._mark_processed(file_path)
                self._sync_vectors(written)

                if "summary" in result:
                    console.print(Panel(
                        f"[bold]{result['summary']}[/bold]\n\n" +
                        "\n".join(f"• {d}" for d in result.get("discoveries", [])),
                        title=f"Done: {file_path.name}",
                        border_style="green"
                    ))

                count = self.state.get("ingest_count", 0)
                if count % AUTO_HOTSPOT_THRESHOLD == 0:
                    self.job_queue.put(("hotspot", None))

        elif job_type == "hotspot":
            result = self.agent.generate_hotspot_analysis()
            self.writer.apply_hotspot(result)
            self._sync_vectors([])
            if "key_themes" in result:
                console.print(Panel(
                    "\n".join(f"• {t}" for t in result.get("key_themes", [])),
                    title="Hotspot themes updated",
                    border_style="magenta"
                ))

    def scan_existing(self):
        """On startup, enqueue files in raw/ that have not yet been processed."""
        unprocessed = []
        for f in self.config.raw_dir.rglob("*"):
            if (f.is_file()
                    and f.suffix.lower() in SUPPORTED_EXTENSIONS
                    and not f.name.startswith('._')
                    and not any(f.match(p) for p in IGNORE_PATTERNS)):
                if not self._already_processed(f):
                    unprocessed.append(f)

        if unprocessed:
            console.print(f"[yellow]Found {len(unprocessed)} unprocessed file(s), queuing...[/yellow]")
            for f in unprocessed:
                self.job_queue.put(("ingest", f))

    def run(self):
        self.running = True
        self.config.pid_file.write_text(str(os.getpid()))

        console.print(Panel(
            f"[bold green]Wiki Watcher started[/bold green]\n"
            f"Watch directory: [cyan]{self.config.raw_dir}[/cyan]\n"
            f"Wiki directory:  [cyan]{self.config.wiki_dir}[/cyan]\n"
            f"Auto hotspot analysis: every {AUTO_HOTSPOT_THRESHOLD} ingests\n"
            f"PID: {os.getpid()}",
            title="LLM Wiki Watcher",
            border_style="blue"
        ))

        self.scan_existing()

        handler = RawDirectoryHandler(self.job_queue, self.config)
        handler.start_debounce_worker()

        observer = Observer()
        observer.schedule(handler, str(self.config.raw_dir), recursive=True)
        observer.start()

        console.print("[green]File watcher started, waiting for new files...[/green]\n")

        def shutdown(sig, frame):
            console.print("\n[yellow]Shutting down...[/yellow]")
            self.running = False
            observer.stop()

        signal.signal(signal.SIGINT, shutdown)
        signal.signal(signal.SIGTERM, shutdown)

        sem = threading.Semaphore(MAX_CONCURRENT_JOBS)

        def _run_job(jt, fp):
            try:
                self.process_job(jt, fp)
            except Exception as e:
                console.print(f"[red]Unhandled exception ({fp}): {e}[/red]")
            finally:
                sem.release()

        while self.running or not self.job_queue.empty():
            try:
                job_type, file_path = self.job_queue.get(timeout=1)
                sem.acquire()
                t = threading.Thread(
                    target=_run_job,
                    args=(job_type, file_path),
                    daemon=True
                )
                t.start()
            except Empty:
                continue

        observer.join()
        if self.config.pid_file.exists():
            self.config.pid_file.unlink()
        console.print("[green]Wiki Watcher stopped.[/green]")


# ── Entry point ───────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="LLM Wiki automation daemon")
    parser.add_argument("--root", default=".", help="path to wiki root directory")
    parser.add_argument("--daemon", action="store_true", help="run in background")
    parser.add_argument("--status", action="store_true", help="show running status")
    parser.add_argument("--stop", action="store_true", help="stop daemon")
    parser.add_argument("--hotspot", action="store_true", help="generate hotspot analysis now")
    parser.add_argument(
        "--model",
        metavar="SPEC",
        help="switch model. Format: 'provider,model' or 'model'. "
             "Examples: anthropic,claude-opus-4-6 | local,Qwen3.6-35B-A3B-FP8 | dashscope,qwen-max"
    )
    args = parser.parse_args()

    wiki_root = Path(args.root).resolve()
    config = WikiConfig(wiki_root)

    if args.model:
        router = ModelRouter(wiki_root)
        router.set_model(args.model)
        if not any([args.daemon, args.hotspot]):
            return

    if args.status:
        router = ModelRouter(wiki_root)
        model_line = f"Active model: [bold]{router.active_provider}[/bold] / [bold]{router.active_model}[/bold]"
        if config.pid_file.exists():
            pid = config.pid_file.read_text().strip()
            state = config.load_state()
            console.print(Panel(
                f"PID: {pid}\n"
                f"Files processed: {len(state.get('processed', {}))}\n"
                f"Total ingests:   {state.get('ingest_count', 0)}\n"
                f"{model_line}",
                title="Wiki Watcher — running",
                border_style="green"
            ))
        else:
            console.print(f"[yellow]Wiki Watcher is not running[/yellow]  |  {model_line}")
        return

    if args.stop:
        if config.pid_file.exists():
            pid = int(config.pid_file.read_text().strip())
            os.kill(pid, signal.SIGTERM)
            console.print(f"[green]Stop signal sent (PID: {pid})[/green]")
        else:
            console.print("[yellow]Wiki Watcher is not running[/yellow]")
        return

    if args.hotspot:
        watcher = WikiWatcher(wiki_root)
        result = watcher.agent.generate_hotspot_analysis()
        watcher.writer.apply_hotspot(result)
        return

    if args.daemon:
        # Use subprocess to start a new process, avoiding fork conflicts with SDK/SSL
        log_path = wiki_root / "wiki_watcher.log"
        cmd = [sys.executable, __file__, "--root", str(wiki_root)]
        log_fd = open(log_path, 'a', buffering=1)  # line-buffered for real-time writes
        proc = subprocess.Popen(cmd, start_new_session=True,
                                stdout=log_fd, stderr=log_fd)
        wiki_root_cfg = WikiConfig(wiki_root)
        wiki_root_cfg.pid_file.write_text(str(proc.pid))
        console.print(f"[green]Wiki Watcher started in background (PID: {proc.pid})[/green]")
        console.print(f"[dim]Log file: {log_path}[/dim]")
        return

    watcher = WikiWatcher(wiki_root)
    watcher.run()


if __name__ == "__main__":
    main()
