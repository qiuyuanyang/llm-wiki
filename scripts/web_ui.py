#!/usr/bin/env python3
# SPDX-License-Identifier: MIT
"""
web_ui.py — LLM Wiki Web Interface

A Flask-based web UI for browsing, searching, and querying the wiki.

Usage:
    python3 scripts/web_ui.py [--port 5000] [--host 0.0.0.0]

Access:
    http://localhost:5000
"""

import json
import re
import sys
import os
import queue
import threading
import subprocess
from pathlib import Path
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_from_directory

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from config import WikiConfig
from vector_store import VectorStore
from embedding_client import EmbeddingClient
from scripts.query_engine import QueryEngine
from scripts.diagnosis_engine import DiagnosisDB, scan_sources
from scripts.vector_ingest import ingest_all as ingest_all_vectors, _clean_orphan_vectors, schedule_orphan_cleanup
from scripts.wiki_watcher import WikiConfig as WatcherConfig, WikiAgent, WikiWriter, ModelRouter

# ── App setup ─────────────────────────────────────────────────────────────────
app = Flask(__name__)
config = WikiConfig()
vector_store = VectorStore(config)
embedding_client = EmbeddingClient(config)
query_engine = QueryEngine(config)
diagnosis_db = DiagnosisDB(config)


# ── Helper functions ──────────────────────────────────────────────────────────


def _format_pdf_text(text: str) -> str:
    """Format raw PDF-extracted text for better Markdown rendering."""
    lines = text.split('\n')
    formatted = []
    for line in lines:
        line = line.strip()
        if not line:
            formatted.append('')
            continue

        # Detect likely section headers:
        # - Lines starting with a number followed by space (e.g. "1 断开...")
        # - Short standalone lines (< 20 chars, no punctuation mid-sentence)
        # - Lines containing "说明/步骤/效果/时间" followed by colon
        is_heading = False

        # Numbered sections: "1 ", "2 ", etc.
        if re.match(r'^\d{1,2}\s+[\u4e00-\u9fff]', line) and len(line) < 80:
            is_heading = True

        # "X." pattern headers
        if re.match(r'^[一二三四五六七八九十]+[、.]\s', line) and len(line) < 80:
            is_heading = True

        # Lines like "测试说明：" "演练步骤：" etc.
        if re.match(r'^[^\n]{1,8}[：:]', line) and len(line.split('：')[0].split(':')[0]) < 20:
            is_heading = True

        # Title-like: single short line at start (e.g. "李朗 11 楼-联通-宝安数据中心演练方案")
        if len(line) < 60 and '-' in line and len(line.split()) < 8:
            is_heading = True

        if is_heading:
            formatted.append(f"### {line}")
        else:
            formatted.append(line)

    return '\n'.join(formatted)


def _format_docx_table(rows) -> str:
    """Format docx Table rows as a Markdown table."""
    if not rows:
        return ''
    cells_list = []
    for row in rows:
        cells = [cell.text.strip() for cell in row.cells]
        cells_list.append(cells)

    max_cols = max(len(c) for c in cells_list) if cells_list else 0
    for cells in cells_list:
        while len(cells) < max_cols:
            cells.append('')

    lines = []
    header = [_escape_md_table_cell(c) for c in cells_list[0]]
    lines.append('| ' + ' | '.join(header) + ' |')
    lines.append('| ' + ' | '.join('---' for _ in cells_list[0]) + ' |')
    for cells in cells_list[1:]:
        escaped = [_escape_md_table_cell(c) for c in cells]
        lines.append('| ' + ' | '.join(escaped) + ' |')
    return '\n'.join(lines)


def _escape_md_table_cell(value: str) -> str:
    """Sanitize a cell value for safe inclusion in a Markdown table row.

    Order matters — backslash MUST be escaped first, otherwise it will
    escape the replacement characters we introduce.

    1. Backslash     \ → \\   (must be first, or it escapes later replacements)
    2. Newlines      \n \r → <br>
    3. Tabs          \t → space
    4. Pipe          |  → \|
    """
    return (value
        .replace('\\', '\\\\')
        .replace('\r\n', '<br>')
        .replace('\n', '<br>')
        .replace('\r', '<br>')
        .replace('\t', ' ')
        .replace('|', '\\|'))


def _format_excel_table(rows: list) -> str:
    """Format Excel rows as a Markdown table."""
    if not rows:
        return ''
    rows = [r for r in rows if any(c is not None and str(c).strip() for c in r)]
    if not rows:
        return ''

    max_cols = max(len(r) for r in rows)
    normalized = []
    for r in rows:
        nr = [str(c) if c is not None else '' for c in r]
        while len(nr) < max_cols:
            nr.append('')
        normalized.append(nr)

    lines = []
    header = [_escape_md_table_cell(c) for c in normalized[0]]
    lines.append('| ' + ' | '.join(header) + ' |')
    lines.append('| ' + ' | '.join('---' for _ in normalized[0]) + ' |')
    for r in normalized[1:]:
        escaped = [_escape_md_table_cell(c) for c in r]
        lines.append('| ' + ' | '.join(escaped) + ' |')
    return '\n'.join(lines)


def _format_pdf_table(table: list) -> str:
    """Format a PDF table as a Markdown table."""
    if not table:
        return ''

    # Find max columns
    max_cols = max(len(row) for row in table)

    # Normalize rows: pad with empty strings, replace None with ''
    rows = []
    for row in table:
        normalized = []
        for c in row:
            val = '' if c is None else str(c).strip()
            normalized.append(val)
        while len(normalized) < max_cols:
            normalized.append('')
        rows.append(normalized)

    # Skip fully empty rows
    rows = [r for r in rows if any(c for c in r)]
    if not rows:
        return ''

    # Build markdown table
    lines = []
    # Header
    header = [_escape_md_table_cell(c) for c in rows[0]]
    lines.append('| ' + ' | '.join(header) + ' |')
    lines.append('| ' + ' | '.join('---' for _ in rows[0]) + ' |')
    # Data rows
    for row in rows[1:]:
        escaped = [_escape_md_table_cell(c) for c in row]
        lines.append('| ' + ' | '.join(escaped) + ' |')

    return '\n'.join(lines)

def read_file_safe(path: Path) -> str:
    """Read file content safely."""
    try:
        return path.read_text(encoding="utf-8")
    except Exception:
        return ""


def strip_frontmatter(content: str) -> str:
    """Remove YAML frontmatter."""
    return re.sub(r'^---\s*\n.*?\n---\s*\n', '', content, flags=re.DOTALL)


def parse_frontmatter(content: str) -> dict:
    """Parse YAML frontmatter.
    Supports: key: value, key: [array], and multi-line list format:
      key:
        - item1
        - item2
    """
    match = re.match(r'^---\s*\n(.*?)\n---\s*\n', content, re.DOTALL)
    if not match:
        return {}
    fm_text = match.group(1)
    fm = {}
    pending_list_key = None
    for line in fm_text.split('\n'):
        stripped = line.strip()
        if not stripped or stripped.startswith('#'):
            continue
        # Check for multi-line list item ("- value" under a known key)
        if pending_list_key is not None and stripped.startswith('- '):
            item = stripped[2:].strip().strip('"').strip("'")
            if item:
                fm[pending_list_key].append(item)
            continue
        # New key: value line (also matches key: with no value for multi-line lists)
        m = re.match(r'^(\w+):\s*(.*)$', stripped)
        if m:
            key, val = m.group(1), m.group(2).strip()
            pending_list_key = None  # reset
            if val.startswith('[') and val.endswith(']'):
                parsed = [v.strip().strip('"').strip("'") for v in val[1:-1].split(',') if v.strip()]
                fm[key] = parsed
            elif (val.startswith('"') and val.endswith('"')) or (val.startswith("'") and val.endswith("'")):
                fm[key] = val[1:-1]
            else:
                # Could be a scalar, or the start of a multi-line list (key: on its own)
                fm[key] = val
                if val == '' or val == '|':
                    # Might be a multi-line list following, track the key
                    pending_list_key = key
                    fm[key] = []
        else:
            # Continuation line or other YAML construct — ignore
            pass
    return fm


def _extract_title(content: str) -> str | None:
    """Extract the first level-1 heading from Markdown body as display title."""
    body = strip_frontmatter(content)
    for line in body.split('\n'):
        m = re.match(r'^#\s+(.+)$', line.strip())
        if m:
            title = m.group(1).strip()
            if title:
                return title
    return None


def get_all_pages() -> list:
    """Get all wiki pages with metadata."""
    pages = []

    # Scan wiki directories
    for dir_name in ['infrastructure', 'sources', 'entities', 'concepts', 'queries', 'comparisons']:
        dir_path = config.wiki_dir / dir_name
        if dir_path.exists():
            for f in dir_path.glob('*.md'):
                if f.name in ('index.md', 'log.md', 'dashboard.md'):
                    continue
                content = read_file_safe(f)
                fm = parse_frontmatter(content)
                title = _extract_title(content)
                pages.append({
                    'path': str(f.relative_to(config.wiki_root)),
                    'name': title if title else f.stem.replace('-', ' ').title(),
                    'type': dir_name,
                    'category': fm.get('category', ''),
                    'tags': fm.get('tags', []),
                    'updated': fm.get('updated', ''),
                    'size': len(content),
                })

    # Scan diagnoses
    diag_dir = config.diagnoses_dir
    if diag_dir.exists():
        for f in diag_dir.glob('*.md'):
            content = read_file_safe(f)
            fm = parse_frontmatter(content)
            title = _extract_title(content)
            pages.append({
                'path': str(f.relative_to(config.wiki_root)),
                'name': title if title else f.stem.replace('-', ' ').title(),
                'type': 'diagnosis',
                'severity': fm.get('severity', 'medium'),
                'status': fm.get('status', 'historical'),
                'tags': fm.get('tags', []),
                'updated': fm.get('updated', ''),
                'size': len(content),
            })

    # Sort by updated date
    pages.sort(key=lambda x: x.get('updated', ''), reverse=True)
    return pages


def get_page_content(page_path: str) -> dict:
    """Get page content with metadata."""
    full_path = config.wiki_root / page_path
    if not full_path.exists():
        return {'error': '页面未找到'}

    content = read_file_safe(full_path)
    fm = parse_frontmatter(content)
    body = strip_frontmatter(content)

    return {
        'path': page_path,
        'name': full_path.stem.replace('-', ' ').title(),
        'type': fm.get('type', 'page'),
        'frontmatter': fm,
        'content': body,
    }


# ── Full ingest pipeline (LLM → Wiki → Vectors) ──────────────────────────────
# Queue-based serial processing to avoid concurrent LLM call conflicts
_ingest_jobs = {}          # filename → {'status': 'running'|'done'|'error', 'message': ...}
_ingest_lock = threading.Lock()
_ingest_queue = queue.Queue()  # serial queue for raw file paths
_ingest_worker_started = False


def _ingest_worker():
    """Background worker that processes ingest jobs one at a time."""
    while True:
        raw_filepath = _ingest_queue.get()
        if raw_filepath is None:  # sentinel to stop
            _ingest_queue.task_done()
            break

        filename = raw_filepath.name
        with _ingest_lock:
            _ingest_jobs[filename] = {'status': 'running', 'message': '正在通过 LLM 处理文档...'}

        try:
            # Step 1: LLM processing
            watcher_cfg = WatcherConfig(config.root)
            agent = WikiAgent(watcher_cfg)
            writer = WikiWriter(watcher_cfg)

            result = agent.ingest_file(raw_filepath)

            if "error" in result:
                with _ingest_lock:
                    _ingest_jobs[filename] = {'status': 'error', 'message': f'LLM 处理失败: {result.get("error", "未知错误")}'}
                _ingest_queue.task_done()
                continue

            # Step 2: Write wiki pages
            written = writer.apply_result(result)

            if not written:
                with _ingest_lock:
                    _ingest_jobs[filename] = {'status': 'error', 'message': '未能写入 wiki 页面'}
                _ingest_queue.task_done()
                continue

            # Step 3: Sync vector store
            ingest_all_vectors(config, embedding_client, show_progress=False)

            summary = result.get('summary', '处理完成')
            with _ingest_lock:
                _ingest_jobs[filename] = {
                    'status': 'done',
                    'message': f'✅ 已向量化并同步索引 — {summary}',
                    'new_pages': result.get('new_pages', []),
                }

        except Exception as e:
            with _ingest_lock:
                _ingest_jobs[filename] = {'status': 'error', 'message': f'处理出错: {str(e)}'}
        finally:
            _ingest_queue.task_done()


def _enqueue_ingest(raw_filepath: Path):
    """Add an ingest job to the serial processing queue."""
    global _ingest_worker_started
    _ingest_queue.put(raw_filepath)

    # Start the worker thread if not already running
    if not _ingest_worker_started:
        _ingest_worker_started = True
        t = threading.Thread(target=_ingest_worker, daemon=True)
        t.start()


@app.route('/api/sources/ingest-status')
def api_ingest_status():
    """API: Check status of ingest jobs."""
    with _ingest_lock:
        return jsonify(dict(_ingest_jobs))


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    """Home page."""
    pages = get_all_pages()
    stats = {
        'total_pages': len(pages),
        'by_type': {},
        'total_vectors': vector_store.count(),
        'diagnoses': diagnosis_db.get_stats()['total'],
    }
    for p in pages:
        t = p['type']
        stats['by_type'][t] = stats['by_type'].get(t, 0) + 1

    return render_template('index.html', pages=pages, stats=stats)


@app.route('/page/<path:page_path>')
def page(page_path):
    """View a wiki page."""
    data = get_page_content(page_path)
    if 'error' in data:
        return render_template('error.html', error=data['error']), 404
    return render_template('page.html', **data)


@app.route('/search')
def search():
    """Search page."""
    query = request.args.get('q', '')
    results = []

    if query:
        # Vector search
        try:
            query_vector = embedding_client.embed_single(query)
            results = vector_store.search(query_vector, top_k=20)
        except Exception as e:
            return render_template('error.html', error=f'搜索出错: {e}'), 500

    return render_template('search.html', query=query, results=results)


@app.route('/query')
def query_page():
    """Intelligent Q&A page."""
    return render_template('query.html')


@app.route('/api/query', methods=['POST'])
def api_query():
    """API: Intelligent Q&A."""
    data = request.json
    question = data.get('question', '')
    if not question:
        return jsonify({'error': '问题不能为空'}), 400

    try:
        result = query_engine.query(question, top_k=8)
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/search', methods=['POST'])
def api_search():
    """API: Vector search."""
    data = request.json
    query = data.get('query', '')
    top_k = data.get('top_k', 20)

    if not query:
        return jsonify({'error': '搜索关键词不能为空'}), 400

    try:
        query_vector = embedding_client.embed_single(query)
        results = vector_store.search(query_vector, top_k=top_k)
        return jsonify(results)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/diagnoses')
def diagnoses_page():
    """Diagnosis knowledge base page."""
    severity = request.args.get('severity', '')
    status = request.args.get('status', '')
    results = diagnosis_db.list_all(severity=severity, status=status, limit=100)
    stats = diagnosis_db.get_stats()
    return render_template('diagnoses.html', results=results, stats=stats)


@app.route('/api/diagnoses/search', methods=['POST'])
def api_diagnoses_search():
    """API: Search diagnoses."""
    data = request.json
    query = data.get('query', '')
    if not query:
        return jsonify({'error': '搜索关键词不能为空'}), 400

    results = diagnosis_db.search(query, limit=50)
    return jsonify(results)


@app.route('/topology')
def topology_page():
    """Topology graph page."""
    return render_template('topology.html')


# ── Source Document Management ────────────────────────────────────────────────

@app.route('/upload')
def upload_page():
    """Source document upload & management page."""
    return render_template('upload.html')


def _find_wiki_source_for_raw(raw_filename: str) -> Path | None:
    """Find the corresponding wiki/sources/ file for a raw source file.

    Checks by raw_file frontmatter field and by stem name.
    Supports .md, .xlsx, .xls, .docx, .doc extensions.
    """
    wiki_src_dir = config.wiki_dir / 'sources'
    if not wiki_src_dir.exists():
        return None

    # Search across all wiki file types (not just .md)
    wiki_files = []
    for ext in ('*.md', '*.xlsx', '*.xls', '*.docx', '*.doc', '*.pdf'):
        wiki_files.extend(wiki_src_dir.glob(ext))

    for wf in wiki_files:
        content = read_file_safe(wf)
        fm = parse_frontmatter(content)
        # Check raw_file frontmatter
        raw_fm = fm.get('raw_file', '')
        if raw_fm and raw_filename in raw_fm:
            return wf
        # Fallback: match stem (normalized)
        raw_stem = Path(raw_filename).stem.lower().replace(' ', '-').replace('_', '-')
        wiki_stem = wf.stem.lower()
        if raw_stem == wiki_stem:
            return wf

    return None


def _source_match_variants(raw_filename: str) -> set:
    """Generate all possible variants of a raw filename that might appear in frontmatter sources."""
    from pathlib import Path
    variants = set()
    stem = Path(raw_filename).stem
    variants.add(raw_filename)                      # 小怡家费用.xlsx
    variants.add(raw_filename + '.md')              # 小怡家费用.xlsx.md
    variants.add(stem)                              # 小怡家费用
    variants.add(stem + '.md')                      # 小怡家费用.md
    variants.add('raw/sources/' + raw_filename)     # raw/sources/小怡家费用.xlsx
    variants.add('wiki/sources/' + raw_filename)    # wiki/sources/小怡家费用.xlsx
    variants.add('wiki/sources/' + raw_filename + '.md')  # wiki/sources/小怡家费用.xlsx.md
    return variants


def _source_matches(sources_val, raw_filename: str) -> bool:
    """Check if a frontmatter 'sources' field references the given raw file.

    Handles various formats:
    - YAML array: ["小怡家费用.xlsx.md"]
    - YAML array with paths: ["raw/sources/小怡家费用.xlsx"]
    - Plain string: "小怡家费用.xlsx"
    - Just stem: ["小怡家费用"]
    """
    if not sources_val:
        return False

    sources_list = sources_val if isinstance(sources_val, list) else [sources_val]

    variants = _source_match_variants(raw_filename)
    stem = None  # lazy init

    for s in sources_list:
        if not isinstance(s, str) or not s.strip():
            continue
        s_lower = s.lower()
        # Check all variants (any direction)
        for v in variants:
            v_lower = v.lower()
            if v_lower in s_lower or s_lower in v_lower:
                return True
        # Also check stem containment
        if stem is None:
            from pathlib import Path as _P
            stem = _P(raw_filename).stem.lower()
        if stem in s_lower:
            return True

    return False


def _find_derived_pages_for_source(raw_filename: str) -> list:
    """Find all wiki pages (infrastructure/concepts/entities/etc.) that reference
    the given raw source file in their 'sources' frontmatter field.
    Returns list of (Path, relative_path) tuples.
    """
    derived = []
    wiki_root = config.wiki_dir
    search_dirs = [wiki_root / 'infrastructure', wiki_root / 'concepts',
                   wiki_root / 'entities', wiki_root / 'queries',
                   wiki_root / 'comparisons', wiki_root / 'sources']

    for d in search_dirs:
        if not d.exists():
            continue
        for ext in ('*.md', '*.xlsx', '*.xls', '*.docx', '*.doc'):
            for wf in d.glob(ext):
                content = read_file_safe(wf)
                fm = parse_frontmatter(content)
                sources = fm.get('sources', [])
                if _source_matches(sources, raw_filename):
                    rel = str(wf.relative_to(config.root))
                    derived.append((wf, rel))
    return derived


def _delete_wiki_and_vector(raw_filename: str) -> list:
    """Delete wiki page(s) and vector entries associated with a raw source file.
    This includes the main source page AND any derivative pages
    (infrastructure/concepts/entities) that reference the source.
    Returns list of actions taken.
    """
    actions = []

    # 1. Delete main wiki source page
    wiki_file = _find_wiki_source_for_raw(raw_filename)
    if wiki_file and wiki_file.exists():
        wiki_rel = str(wiki_file.relative_to(config.root))
        wiki_file.unlink()
        vector_store.delete_vector(wiki_rel)
        actions.append(f'已删除 wiki 页面: {wiki_file.name}')
        actions.append(f'已删除向量: {wiki_rel}')
    else:
        # Fallback: try common page_path patterns for vectors
        raw_stem = Path(raw_filename).stem.replace(' ', '-').replace('_', '-').lower()
        if not raw_stem.endswith('.md'):
            raw_stem = raw_stem + '.md'
        wiki_path = f'wiki/sources/{raw_stem}'
        if vector_store.delete_vector(wiki_path):
            actions.append(f'已删除向量: {wiki_path}')

    # 2. Find and delete ALL derivative pages that reference this source
    derived = _find_derived_pages_for_source(raw_filename)
    for wf, rel in derived:
        if wf.exists():
            wf.unlink()
            vector_store.delete_vector(rel)
            actions.append(f'已删除衍生页面: {wf.name} ({rel})')

    return actions


@app.route('/api/sources')
def api_list_sources():
    """API: List all uploaded source files with optional text filter + pagination."""
    raw_dir = config.raw_dir
    if not raw_dir.exists():
        return jsonify({'files': [], 'total': 0, 'page': 1, 'per_page': 10, 'total_pages': 0})

    # Build wiki map
    wiki_src_dir = config.wiki_dir / 'sources'
    raw_to_wiki = {}
    if wiki_src_dir.exists():
        # Scan all files (not just .md) — docx/xlsx sources keep original ext
        for wf in list(wiki_src_dir.glob('*.md')) + list(wiki_src_dir.glob('*.xlsx')) + list(wiki_src_dir.glob('*.docx')) + list(wiki_src_dir.glob('*.doc')):
            content = read_file_safe(wf)
            fm = parse_frontmatter(content)
            raw_fm = fm.get('raw_file', '')
            if raw_fm:
                raw_name = Path(raw_fm).name
                raw_to_wiki[raw_name] = wf.name

    # Merge ingest job status into file list
    with _ingest_lock:
        job_status = dict(_ingest_jobs)  # filename → {status, message}

    files = []
    for f in sorted(raw_dir.iterdir()):
        if f.is_file() and not f.name.startswith('.'):
            stat = f.stat()
            raw_name = f.name
            wiki_name = raw_to_wiki.get(raw_name)

            # Determine processing state from ingest jobs
            job = job_status.get(raw_name)
            is_processing = job is not None and job.get('status') == 'running'
            has_errored = job is not None and job.get('status') == 'error'

            # Optional text filter (filename + wiki content snippet)
            keyword = request.args.get('keyword', '').strip().lower()
            if keyword:
                match = keyword in raw_name.lower()
                if not match and wiki_name:
                    wiki_path = wiki_src_dir / wiki_name
                    wiki_content = read_file_safe(wiki_path).lower()
                    match = keyword in wiki_content
                if not match:
                    continue

                    # If watcher daemon processed it (wiki page exists) but Web UI still shows error,
            # override the error state — watcher is the authoritative source.
            if has_errored and wiki_name:
                has_errored = False
                job = None

            files.append({
                'name': raw_name,
                'size': stat.st_size,
                'mtime': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M'),
                'ingested': wiki_name is not None and not is_processing,
                'processing': is_processing,
                'processing_message': job.get('message') if is_processing else None,
                'error': job.get('message') if has_errored else None,
                'wiki_file': wiki_name,
            })

    # Sort by mtime descending
    files.sort(key=lambda x: x['mtime'], reverse=True)

    # Pagination
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 10, type=int)
    per_page = min(per_page, 50)
    total = len(files)
    total_pages = (total + per_page - 1) // per_page if total else 0
    start = (page - 1) * per_page
    page_files = files[start:start + per_page]

    return jsonify({
        'files': page_files,
        'total': total,
        'page': page,
        'per_page': per_page,
        'total_pages': total_pages,
    })


@app.route('/api/sources/search', methods=['POST'])
def api_search_sources():
    """API: Semantic search across source document content via vector store."""
    data = request.json
    query = data.get('query', '')
    page = data.get('page', 1)
    per_page = data.get('per_page', 10)
    per_page = min(per_page, 50)

    if not query:
        return jsonify({'error': '搜索关键词不能为空'}), 400

    try:
        # Search vectors filtered by source type
        query_vector = embedding_client.embed_single(query)
        results = vector_store.search(query_vector, top_k=200, page_type='source')

        # Enrich with raw file info
        raw_dir = config.raw_dir
        wiki_src_dir = config.wiki_dir / 'sources'
        wiki_to_raw = {}  # wiki_filename -> raw_filename
        if wiki_src_dir.exists():
            for wf in wiki_src_dir.glob('*.md'):
                content = read_file_safe(wf)
                fm = parse_frontmatter(content)
                raw_fm = fm.get('raw_file', '')
                if raw_fm:
                    wiki_to_raw[wf.name] = Path(raw_fm).name

        enriched = []
        for r in results:
            wiki_name = Path(r['page_path']).name
            raw_name = wiki_to_raw.get(wiki_name)
            raw_info = {}
            if raw_name:
                raw_path = raw_dir / raw_name
                if raw_path.exists():
                    stat = raw_path.stat()
                    raw_info = {
                        'size': stat.st_size,
                        'mtime': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M'),
                    }
            enriched.append({
                'name': raw_name or wiki_name,
                'wiki_file': wiki_name,
                'page_path': r['page_path'],
                'title': r['title'],
                'score': r['score'],
                **raw_info,
            })

        # Paginate
        total = len(enriched)
        total_pages = (total + per_page - 1) // per_page if total else 0
        start = (page - 1) * per_page
        page_results = enriched[start:start + per_page]

        return jsonify({
            'files': page_results,
            'total': total,
            'page': page,
            'per_page': per_page,
            'total_pages': total_pages,
            'query': query,
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/sources/upload', methods=['POST'])
def api_upload_source():
    """API: Upload a source file and trigger full ingest pipeline."""
    if 'file' not in request.files:
        return jsonify({'error': '没有文件'}), 400

    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    # Sanitize filename
    filename = re.sub(r'[^\w\u4e00-\u9fff.\-]', '_', file.filename)
    if not filename:
        return jsonify({'error': '无效文件名'}), 400

    # Ensure raw/sources dir exists
    raw_dir = config.raw_dir
    raw_dir.mkdir(parents=True, exist_ok=True)

    dest = raw_dir / filename

    # Prevent overwriting - add suffix if exists
    if dest.exists():
        stem = dest.stem
        suffix = dest.suffix
        counter = 1
        while dest.exists():
            dest = raw_dir / f"{stem}_{counter}{suffix}"
            counter += 1
        filename = dest.name

    file.save(str(dest))

    # Enqueue for serial processing (avoid concurrent LLM conflicts)
    _enqueue_ingest(dest)

    return jsonify({
        'message': f'已上传: {filename}，正在后台处理...',
        'filename': filename,
        'ingest_started': True,
    })


@app.route('/api/sources/<path:filename>/content')
def api_source_content(filename):
    """API: Get source file content for preview."""
    raw_dir = config.raw_dir
    filepath = raw_dir / filename

    # Prevent path traversal
    filepath = filepath.resolve()
    if not str(filepath).startswith(str(raw_dir.resolve())):
        return jsonify({'error': '非法路径'}), 403

    if not filepath.exists():
        return jsonify({'error': '文件不存在'}), 404

    try:
        ext = filepath.suffix.lower()

        if ext == '.pdf':
            # Extract text from PDF using pdfplumber with formatting
            try:
                import pdfplumber
            except ImportError:
                return jsonify({'error': 'PDF 提取需要安装 pdfplumber：pip install pdfplumber'}), 500

            formatted = []
            with pdfplumber.open(str(filepath)) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    formatted.append(f"--- 第 {i} 页 ---")
                    formatted.append("")

                    tables = page.extract_tables()
                    has_table = False
                    for table in tables:
                        if table:
                            formatted.append(_format_pdf_table(table))
                            formatted.append("")
                            has_table = True

                    # If page has a table, skip plain text to avoid duplication
                    if has_table:
                        formatted.append("")
                        continue

                    page_text = page.extract_text()
                    if page_text:
                        formatted.append(_format_pdf_text(page_text))
                        formatted.append("")
                    else:
                        formatted.append("（此页无文本内容，可能是扫描件/图片）")
                        formatted.append("")

            content = "\n".join(formatted)
            return jsonify({'content': content, 'filename': filename, 'type': 'pdf'})

        if ext == '.docx':
            try:
                from docx import Document
                doc = Document(str(filepath))
                parts = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        parts.append(p.text.strip())
                for table in doc.tables:
                    rows = table.rows
                    if rows:
                        parts.append(_format_docx_table(rows))
                content = "\n\n".join(parts)
                return jsonify({'content': content, 'filename': filename, 'type': 'docx'})
            except ImportError:
                return jsonify({'error': 'DOCX 提取需要安装 python-docx：pip install python-docx'}), 500

        if ext == '.doc':
            try:
                import subprocess
                result = subprocess.run(
                    ['antiword', str(filepath)],
                    capture_output=True, text=True, timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return jsonify({'content': result.stdout, 'filename': filename, 'type': 'doc'})
                return jsonify({'error': 'antiword 提取失败，请安装: apt install antiword'}), 500
            except FileNotFoundError:
                return jsonify({'error': '.doc 格式需要安装 antiword：apt install antiword。建议转为 .docx'}), 500

        if ext in ('.xlsx', '.xls'):
            try:
                if ext == '.xlsx':
                    import openpyxl
                    wb = openpyxl.load_workbook(str(filepath), data_only=True)
                    sheets_data = []
                    for sheet_name in wb.sheetnames:
                        ws = wb[sheet_name]
                        rows = list(ws.iter_rows(values_only=True))
                        sheets_data.append((sheet_name, rows))
                else:
                    import xlrd
                    wb = xlrd.open_workbook(str(filepath))
                    sheets_data = []
                    for si in range(wb.nsheets):
                        ws = wb.sheet_by_index(si)
                        rows = []
                        for r in range(ws.nrows):
                            rows.append([ws.cell_value(r, c) for c in range(ws.ncols)])
                        sheets_data.append((ws.name, rows))

                parts = []
                for sheet_name, rows in sheets_data:
                    parts.append(f"## 工作表: {sheet_name}")
                    rows_clean = [r for r in rows if any(c is not None and str(c).strip() for c in r)]
                    if rows_clean:
                        parts.append(_format_excel_table(rows_clean))
                    else:
                        parts.append("_(空工作表)_")
                content = "\n\n".join(parts)
                return jsonify({'content': content, 'filename': filename, 'type': 'excel'})
            except ImportError:
                return jsonify({'error': 'Excel 提取需要安装 openpyxl/xlrd：pip install openpyxl xlrd'}), 500

        else:
            content = filepath.read_text(encoding='utf-8')
            return jsonify({'content': content, 'filename': filename})
    except UnicodeDecodeError:
        # Try other encodings for non-PDF files
        for enc in ['gbk', 'gb2312', 'latin-1']:
            try:
                content = filepath.read_text(encoding=enc)
                return jsonify({'content': content, 'filename': filename})
            except UnicodeDecodeError:
                continue
        return jsonify({'error': '无法解码文件，可能包含二进制数据'}), 500
    except Exception as e:
        return jsonify({'error': f'读取失败: {e}'}), 500


@app.route('/api/sources/<path:filename>/download')
def api_download_source(filename):
    """API: Download a source file."""
    raw_dir = config.raw_dir
    filepath = raw_dir / filename

    filepath = filepath.resolve()
    if not str(filepath).startswith(str(raw_dir.resolve())):
        return jsonify({'error': '非法路径'}), 403

    if not filepath.exists():
        return jsonify({'error': '文件不存在'}), 404

    from flask import send_file
    return send_file(
        str(filepath),
        as_attachment=True,
        download_name=filename,
    )


@app.route('/api/sources/<path:filename>', methods=['DELETE'])
def api_delete_source(filename):
    """API: Delete a source file and clean up wiki + vectors."""
    raw_dir = config.raw_dir
    filepath = raw_dir / filename

    # Prevent path traversal
    filepath = filepath.resolve()
    if not str(filepath).startswith(str(raw_dir.resolve())):
        return jsonify({'error': '非法路径'}), 403

    if not filepath.exists():
        return jsonify({'error': '文件不存在'}), 404

    try:
        # Step 1: Delete corresponding wiki page and vector entries
        cleanup_actions = _delete_wiki_and_vector(filename)

        # Step 2: Delete raw file
        filepath.unlink()

        # Step 3: Re-sync vector store to ensure consistency (async, non-blocking)
        schedule_orphan_cleanup(config, embedding_client, show_progress=False)

        messages = [f'已删除源文件: {filename}'] + cleanup_actions
        return jsonify({
            'message': '删除完成',
            'details': messages,
        })
    except Exception as e:
        return jsonify({'error': f'删除失败: {e}'}), 500


@app.route('/api/sources/<path:filename>/ingest', methods=['POST'])
def api_ingest_source(filename):
    """API: Trigger ingest for a specific source file."""
    raw_dir = config.raw_dir
    filepath = raw_dir / filename

    filepath = filepath.resolve()
    if not str(filepath).startswith(str(raw_dir.resolve())):
        return jsonify({'error': '非法路径'}), 403

    if not filepath.exists():
        return jsonify({'error': '文件不存在'}), 404

    # Enqueue for serial processing
    _enqueue_ingest(filepath)

    return jsonify({
        'message': f'正在处理: {filename}',
        'filename': filename,
        'ingest_started': True,
    })


@app.route('/api/sources/<path:filename>/reprocess', methods=['POST'])
def api_reprocess_source(filename):
    """API: Reprocess a source file (delete wiki entry first, then re-ingest)."""
    raw_dir = config.raw_dir
    filepath = raw_dir / filename

    filepath = filepath.resolve()
    if not str(filepath).startswith(str(raw_dir.resolve())):
        return jsonify({'error': '非法路径'}), 403

    if not filepath.exists():
        return jsonify({'error': '文件不存在'}), 404

    try:
        # Clean up existing wiki + vector entries
        _delete_wiki_and_vector(filename)

        # Re-ingest via serial queue
        _enqueue_ingest(filepath)

        return jsonify({
            'message': f'正在重新处理: {filename}',
            'filename': filename,
            'ingest_started': True,
        })
    except Exception as e:
        return jsonify({'error': f'重新处理失败: {e}'}), 500


@app.route('/api/vectors/cleanup-orphan', methods=['POST'])
def api_cleanup_orphan_vectors():
    """API: Remove vector entries whose wiki files no longer exist on disk."""
    try:
        orphaned = _clean_orphan_vectors(config, show_progress=False)
        return jsonify({
            'message': f'清理完成，共删除 {len(orphaned)} 个孤儿向量条目',
            'orphaned': orphaned,
            'count': len(orphaned),
        })
    except Exception as e:
        return jsonify({'error': f'清理失败: {e}'}), 500


@app.route('/api/stats')
def api_stats():
    """API: Wiki statistics."""
    pages = get_all_pages()
    stats = {
        'total_pages': len(pages),
        'by_type': {},
        'total_vectors': vector_store.count(),
        'vector_stats': vector_store.get_stats(),
        'diagnoses': diagnosis_db.get_stats(),
    }
    for p in pages:
        t = p['type']
        stats['by_type'][t] = stats['by_type'].get(t, 0) + 1
    return jsonify(stats)


@app.route('/topology.html')
def serve_topology():
    """Serve the generated topology HTML."""
    topology_path = config.wiki_dir / 'topology.html'
    if topology_path.exists():
        return send_from_directory(str(config.wiki_dir), 'topology.html')
    return 'Topology not generated. Run: ./wiki.sh graph', 404


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    import argparse
    parser = argparse.ArgumentParser(description='LLM Wiki Web UI')
    parser.add_argument('--port', type=int, default=5000, help='Port to listen on (default: 5000)')
    parser.add_argument('--host', default='0.0.0.0', help='Host to bind (default: 0.0.0.0)')
    parser.add_argument('--debug', action='store_true', help='Enable debug mode')
    args = parser.parse_args()

    print(f"🌐 LLM Wiki Web UI starting...")
    print(f"   Access: http://localhost:{args.port}")
    print(f"   Press Ctrl+C to stop")

    app.run(host=args.host, port=args.port, debug=args.debug)


if __name__ == '__main__':
    main()
